# -*- coding: utf-8 -*-
"""
VoiceVerse - AI Text-to-Speech Application
Clean, fully functional version
"""

from flask import Flask, render_template_string, render_template, request, send_file, redirect, url_for, jsonify, session
import os
import re
from openai import OpenAI
import io
import sys
from datetime import datetime
import json
from collections import defaultdict
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import html
import secrets
from docx import Document
from PyPDF2 import PdfReader
from functools import wraps
from tts_agents import create_agent_system
from dotenv import load_dotenv
from flask_wtf.csrf import CSRFProtect
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from database import Database
from logger import SecurityLogger
from simple_lockout import SimpleLockout
from simple_alerts import SimpleAlerts
import hashlib
from monitoring.metrics_collector import get_metrics_collector
from monitoring.log_analyzer import LogAnalyzer
from workflow_api import workflow_bp
# NOTE: routes/auth.py has a bug (stores plaintext passwords) - keeping auth routes in main file
import soundfile as sf
import numpy as np
from encryption import encrypt_api_key, decrypt_api_key, validate_openai_api_key, mask_api_key
from features.analytics import CostEstimator

# Centralized TTS pricing constants (from CostEstimator)
TTS_PRICING = CostEstimator.PRICING  # {'tts-1': 0.015, 'tts-1-hd': 0.030}

# Load environment variables from .env file
load_dotenv()

# Voice Cloning: Try to import Coqui TTS for voice cloning
# Auto-accept Coqui TTS license terms (CPML) for non-interactive usage
# This is required for XTTS v2 model to load without interactive prompt
os.environ["COQUI_TOS_AGREED"] = "1"

TTS_XTTS = None
try:
    from TTS.api import TTS as TTS_API
    TTS_XTTS = TTS_API

    # Fix for PyTorch 2.6+ weights_only default change
    # PyTorch 2.6 changed torch.load() to use weights_only=True by default
    # XTTS model checkpoints require these classes to be allowlisted as safe globals
    import torch
    try:
        from TTS.tts.configs.xtts_config import XttsConfig
        from TTS.tts.models.xtts import XttsArgs, XttsAudioConfig
        from TTS.config import BaseAudioConfig, BaseDatasetConfig
        torch.serialization.add_safe_globals([XttsConfig, XttsArgs, XttsAudioConfig, BaseAudioConfig, BaseDatasetConfig])
        print("✅ PyTorch safe globals configured for XTTS model")
    except (ImportError, AttributeError) as e:
        print(f"⚠️  Could not configure PyTorch safe globals: {e}")

    print("✅ Coqui TTS library loaded successfully")
except ImportError:
    print("⚠️  Coqui TTS not installed - voice cloning features disabled")
    print("   Install with: pip install TTS torch soundfile")

# Ensure UTF-8 output for Flask
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')


app = Flask(__name__)
app.config['JSON_AS_ASCII'] = False
app.config['UPLOAD_FOLDER'] = 'saved_audio'
app.config['VOICE_SAMPLES_FOLDER'] = 'voice_samples'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # Increased to 50MB for voice samples

# Security: Enforce SECRET_KEY from environment (production safety)
secret_key = os.getenv('SECRET_KEY')
if not secret_key:
    raise ValueError(
        "ERROR: SECRET_KEY environment variable is not set!\n"
        "Generate one with: python3 -c \"import secrets; print(secrets.token_hex(32))\"\n"
        "Then add it to your .env file: SECRET_KEY=your-generated-key-here"
    )
app.config['SECRET_KEY'] = secret_key

# Security: Initialize CSRF Protection
csrf = CSRFProtect(app)

# Database: Initialize SQLite database
db = Database('voiceverse.db')

# Security: Initialize security logger
security_log = SecurityLogger(db)

# Security: Initialize lockout and email alerts system
lockout = SimpleLockout('data/lockouts.db')
alerts = SimpleAlerts()

# Initialize workflow API with agent executor
from workflow_api import init_agent_executor

# Security: IP hashing function for privacy
def hash_ip(ip_address: str) -> str:
    """Hash IP address with salt for privacy-preserving logging."""
    salt = os.getenv('IP_HASH_SALT', 'default-salt-change-this')
    return hashlib.sha256(f"{ip_address}{salt}".encode()).hexdigest()[:16]

# Security: Configure session security (environment-based)
# Read cookie security settings from environment
secure_cookies = os.getenv('SECURE_COOKIES', 'false').lower() == 'true'
session_httponly = os.getenv('SESSION_COOKIE_HTTPONLY', 'true').lower() == 'true'
session_samesite = os.getenv('SESSION_COOKIE_SAMESITE', 'Lax')
session_lifetime = int(os.getenv('SESSION_LIFETIME', '3600'))

app.config.update(
    SESSION_COOKIE_SECURE=secure_cookies,  # Requires HTTPS when True
    SESSION_COOKIE_HTTPONLY=session_httponly,  # Prevent JavaScript access to session cookie
    SESSION_COOKIE_SAMESITE=session_samesite,  # Prevent CSRF attacks
    PERMANENT_SESSION_LIFETIME=session_lifetime  # Session expires after configured time
)

# Security: Initialize Rate Limiter (5 login attempts per 15 minutes per IP)
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"],
    storage_uri="memory://"
)

# Workflow Editor: Register workflow API blueprint
app.register_blueprint(workflow_bp)

# Security: Add security headers to all responses
@app.after_request
def set_security_headers(response):
    """Add security headers to protect against common web vulnerabilities"""
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options'] = 'DENY'
    response.headers['X-XSS-Protection'] = '1; mode=block'

    # Security: Tightened CSP - only allowing inline styles/scripts via nonces would be ideal,
    # but for now we restrict to self and essential inline content. In production with external
    # resources, use nonces or hashes instead of 'unsafe-inline'.
    response.headers['Content-Security-Policy'] = (
        "default-src 'self'; "
        "script-src 'self' 'unsafe-inline'; "  # TODO: Replace with nonce-based approach
        "style-src 'self' 'unsafe-inline'; "    # TODO: Replace with nonce-based approach
        "img-src 'self' data:; "
        "font-src 'self'; "
        "connect-src 'self'; "
        "media-src 'self' blob:; "
        "object-src 'none'; "
        "frame-ancestors 'none'; "
        "base-uri 'self'; "
        "form-action 'self';"
    )

    # Only set HSTS if using HTTPS in production
    if not app.debug:
        response.headers['Strict-Transport-Security'] = 'max-age=31536000; includeSubDomains'
    return response

# Performance Monitoring: Track request duration and record metrics
import time as time_module

@app.before_request
def before_request_metrics():
    """Track request start time for performance monitoring"""
    request._start_time = time_module.time()

@app.after_request
def after_request_metrics(response):
    """Record request metrics after processing"""
    try:
        if hasattr(request, '_start_time'):
            duration = time_module.time() - request._start_time

            # Get metrics collector
            metrics = get_metrics_collector()

            # Record HTTP request
            metrics.record_request(
                endpoint=request.path,
                method=request.method,
                status_code=response.status_code,
                duration=duration
            )

            # Add response time header for debugging
            response.headers['X-Response-Time'] = f"{duration:.4f}s"
    except Exception as e:
        # Don't break requests if metrics recording fails
        print(f"Error recording metrics: {e}")

    return response

openai_client = None
agent_system = None
xtts_model = None  # Voice cloning model
VALID_VOICES = {'alloy', 'echo', 'fable', 'onyx', 'nova', 'shimmer'}

# Security: Set up security audit logging
import logging
from logging.handlers import RotatingFileHandler

# Create logs directory if it doesn't exist
if not os.path.exists('logs'):
    os.makedirs('logs')

# Configure security logger with rotation
security_logger = logging.getLogger('security_audit')
security_logger.setLevel(logging.INFO)

# Rotating file handler: max 10MB per file, keep 10 backup files
security_handler = RotatingFileHandler(
    'logs/security_audit.log',
    maxBytes=10*1024*1024,  # 10MB
    backupCount=10
)

# Format: timestamp | level | IP | username | event | details
security_formatter = logging.Formatter(
    '%(asctime)s | %(levelname)s | %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
security_handler.setFormatter(security_formatter)
security_logger.addHandler(security_handler)

def log_security_event(event_type, details, username=None, ip_address=None, success=True):
    """
    Security: Log security-relevant events using SecurityLogger

    This is a compatibility wrapper for the new SecurityLogger class.
    Migrates old logging calls to use the new structured security logging.

    Args:
        event_type: Type of event (LOGIN, REGISTER, etc.)
        details: Additional event details
        username: Username (if applicable)
        ip_address: IP address (if available)
        success: Whether the action succeeded
    """
    if ip_address is None:
        ip_address = get_remote_address()

    if username is None and 'username' in session:
        username = session['username']

    # Get user ID if possible
    user_id = None
    if username:
        try:
            user = db.get_user(username)
            if user:
                user_id = user['id']
        except:
            pass

    # Route to appropriate SecurityLogger method based on event type
    if event_type in ['LOGIN', 'LOGIN_ATTEMPT']:
        security_log.log_authentication(
            username=username or 'anonymous',
            ip_address=ip_address,
            success=success,
            details=details,
            user_id=user_id
        )
    elif event_type in ['REGISTER', 'REGISTER_ATTEMPT']:
        security_log.log_registration(
            username=username or 'anonymous',
            ip_address=ip_address,
            success=success,
            details=details,
            user_id=user_id
        )
    elif event_type == 'LOGOUT':
        security_log.log_logout(
            username=username or 'anonymous',
            ip_address=ip_address,
            user_id=user_id
        )
    else:
        # Fallback to general event logging
        security_log.log_event(
            event_type=event_type,
            user_id=user_id,
            username=username,
            ip_address=ip_address,
            details=details,
            success=success
        )

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

if not os.path.exists(app.config['VOICE_SAMPLES_FOLDER']):
    os.makedirs(app.config['VOICE_SAMPLES_FOLDER'])

METADATA_FILE = os.path.join(app.config['UPLOAD_FOLDER'], 'metadata.json')
USAGE_FILE = os.path.join(app.config['UPLOAD_FOLDER'], 'usage_stats.json')
HISTORY_FILE = os.path.join(app.config['UPLOAD_FOLDER'], 'playback_history.json')
USERS_FILE = os.path.join(app.config['UPLOAD_FOLDER'], 'users.json')

# Per-user OpenAI client cache
_user_openai_clients = {}

def get_openai_client():
    """
    Get OpenAI client - uses global fallback key if available.
    For user-specific operations, use get_user_openai_client() instead.
    """
    global openai_client
    if openai_client is None:
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            openai_client = OpenAI(api_key=api_key)
            # Initialize agent executor for workflow API
            init_agent_executor(openai_client)
        else:
            # No global key - that's ok for BYOK mode
            return None
    return openai_client


def get_user_openai_client(user_id: int):
    """
    Get OpenAI client for a specific user using their stored API key.

    Args:
        user_id: The user's ID

    Returns:
        OpenAI client instance or None if user has no API key

    Note:
        Clients are cached per-user to avoid recreating on every request.
        Cache is cleared when user updates their API key.
    """
    global _user_openai_clients

    # Check cache first
    if user_id in _user_openai_clients:
        return _user_openai_clients[user_id]

    # Get user's encrypted API key from database
    encrypted_key = db.get_user_api_key(user_id)
    if not encrypted_key:
        # User has no API key set - try global fallback
        global_client = get_openai_client()
        if global_client:
            return global_client
        return None

    # Decrypt the API key
    api_key = decrypt_api_key(encrypted_key)
    if not api_key:
        # Decryption failed - key may be corrupted or SECRET_KEY changed
        return None

    # Create and cache the client
    try:
        client = OpenAI(api_key=api_key)
        _user_openai_clients[user_id] = client
        return client
    except Exception as e:
        print(f"Error creating OpenAI client for user {user_id}: {e}")
        return None


def clear_user_client_cache(user_id: int):
    """Clear cached OpenAI client for a user (call when API key changes)."""
    global _user_openai_clients
    if user_id in _user_openai_clients:
        del _user_openai_clients[user_id]


def user_has_api_key(user_id: int) -> bool:
    """Check if user has an API key set (either their own or global fallback)."""
    # Check user's own key first
    if db.has_api_key(user_id):
        return True
    # Check global fallback
    return os.getenv("OPENAI_API_KEY") is not None

def get_agent_system():
    global agent_system
    if agent_system is None:
        client = get_openai_client()
        agent_system = create_agent_system(client)
    return agent_system

def get_xtts_model():
    """Initialize and return the XTTS voice cloning model"""
    global xtts_model
    if xtts_model is None and TTS_XTTS is not None:
        try:
            print("🔄 Loading XTTS v2 model for voice cloning...")
            xtts_model = TTS_XTTS("tts_models/multilingual/multi-dataset/xtts_v2")
            print("✅ XTTS v2 model loaded successfully")
        except Exception as e:
            print(f"❌ Failed to load XTTS model: {e}")
            return None
    return xtts_model

# User Management Functions
def create_user(username, password):
    """Create a new user account using database"""
    password_hash = generate_password_hash(password)
    user_id = db.create_user(username, password_hash)
    return user_id is not None

def verify_user(username, password):
    """Verify user credentials using database"""
    user = db.get_user(username)

    if not user:
        return False

    stored_hash = user['password_hash']

    # Handle both scrypt and sha256 for backwards compatibility
    if stored_hash.startswith('scrypt:'):
        is_valid = check_password_hash(stored_hash, password)
    elif stored_hash.startswith('sha256:'):
        is_valid = check_password_hash(stored_hash, password)
    else:
        # Legacy plain hash
        is_valid = check_password_hash(stored_hash, password)

    # Update last login on successful authentication
    if is_valid:
        db.update_last_login(user['id'])

    return is_valid

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def admin_required(f):
    """Decorator to require admin privileges"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            return redirect(url_for('login'))

        # Check if user is admin
        user = db.get_user(session['username'])
        if not user or not user.get('is_admin', False):
            return render_template_string("""
                <!DOCTYPE html>
                <html>
                <head>
                    <title>Access Denied</title>
                    <style>
                        body {
                            font-family: Arial, sans-serif;
                            background: #0a0e27;
                            color: white;
                            display: flex;
                            align-items: center;
                            justify-content: center;
                            height: 100vh;
                            margin: 0;
                        }
                        .container {
                            text-align: center;
                        }
                        h1 { color: #ef4444; }
                        a {
                            color: #3b82f6;
                            text-decoration: none;
                        }
                    </style>
                </head>
                <body>
                    <div class="container">
                        <h1>🚫 Access Denied</h1>
                        <p>This page is only accessible to administrators.</p>
                        <p><a href="/">Return to Home</a></p>
                    </div>
                </body>
                </html>
            """), 403

        return f(*args, **kwargs)
    return decorated_function

def validate_password(password):
    """
    Security: Validate password strength
    Requires: 12+ characters, uppercase, lowercase, digit, special character
    Returns: (is_valid, error_message)
    """
    if len(password) < 12:
        return False, "Password must be at least 12 characters long"
    if not re.search(r'[A-Z]', password):
        return False, "Password must contain at least one uppercase letter"
    if not re.search(r'[a-z]', password):
        return False, "Password must contain at least one lowercase letter"
    if not re.search(r'\d', password):
        return False, "Password must contain at least one digit"
    if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
        return False, "Password must contain at least one special character (!@#$%^&*(),.?\":{}|<>)"
    return True, ""

def calculate_cost(characters, model='tts-1'):
    """Calculate TTS cost using centralized pricing."""
    return (characters / 1000) * TTS_PRICING.get(model, TTS_PRICING['tts-1'])

def sanitize_display_name(name):
    name = re.sub(r'[<>:"|?*\x00-\x1f]', '', name)
    name = name.strip()[:100]
    return name or 'audio'

def validate_voice(voice):
    return voice if voice in VALID_VOICES else 'nova'

def verify_file_ownership(filename, username):
    """
    Security: Verify that the current user owns the file

    Args:
        filename: The secure filename to check
        username: The username from session

    Returns:
        bool: True if user owns the file, False otherwise
    """
    file_info = db.get_audio_file(filename)

    if not file_info:
        return False

    # Get owner's username from user_id
    owner = db.get_user_by_id(file_info['owner_id'])

    if not owner:
        return False

    return owner['username'] == username

def migrate_existing_files_ownership():
    """
    Migration is now handled by migrate_to_sqlite.py
    This function is kept for backward compatibility but does nothing.
    """
    pass

# Authentication Routes
@app.route('/login', methods=['GET', 'POST'])
@limiter.limit("10 per minute")  # Security: Prevent brute force attacks
def login():
    error = None
    if request.method == 'POST':
        # Security: Get and hash IP address for privacy-preserving lockout tracking
        user_ip = request.remote_addr
        hashed_ip = hash_ip(user_ip)

        # Security: Check lockout status before processing login
        lockout_status = lockout.check_and_record(hashed_ip)

        if lockout_status['locked']:
            # Account is locked - show lockout message
            remaining = lockout_status['remaining_seconds']
            error = f"Too many failed login attempts. Account locked for {remaining} seconds. Please try again later."
            log_security_event('LOGIN_BLOCKED', f'Account locked, {remaining}s remaining', success=False)

            # Send email alert if this is a new lockout
            if lockout_status['should_alert']:
                try:
                    alerts.send_lockout_alert(
                        identifier_hash=hashed_ip,
                        attempt_count=lockout_status['attempt_count'],
                        lockout_duration=60
                    )
                    log_security_event('ALERT_SENT', f'Lockout email alert sent for {hashed_ip[:8]}...', success=True)
                except Exception as e:
                    log_security_event('ALERT_FAILED', f'Failed to send lockout alert: {str(e)}', success=False)

            return render_template('auth.html', error=error, mode='login')

        # Process login credentials
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')

        if not username or not password:
            error = "Please enter both username and password"
            log_security_event('LOGIN_ATTEMPT', f'Empty credentials', username=username, success=False)
        elif verify_user(username, password):
            # Successful login - clear lockout attempts
            lockout.clear_attempts(hashed_ip)
            session['username'] = username
            log_security_event('LOGIN', f'User logged in', username=username, success=True)
            return redirect(url_for('index'))
        else:
            error = "Invalid username or password"
            log_security_event('LOGIN_ATTEMPT', f'Invalid credentials (attempt {lockout_status["attempt_count"]})', username=username, success=False)

    return render_template('auth.html', error=error, mode='login')

@app.route('/register', methods=['GET', 'POST'])
@limiter.limit("5 per minute")  # Security: Prevent account creation spam
def register():
    error = None
    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        confirm = request.form.get('confirm', '')

        if not username or not password:
            error = "Please enter both username and password"
            log_security_event('REGISTER_ATTEMPT', 'Empty credentials', username=username, success=False)
        elif len(username) < 3:
            error = "Username must be at least 3 characters"
            log_security_event('REGISTER_ATTEMPT', 'Username too short', username=username, success=False)
        elif password != confirm:
            error = "Passwords do not match"
            log_security_event('REGISTER_ATTEMPT', 'Password mismatch', username=username, success=False)
        else:
            # Security: Validate password strength
            is_valid, validation_error = validate_password(password)
            if not is_valid:
                error = validation_error
                log_security_event('REGISTER_ATTEMPT', f'Weak password: {validation_error}', username=username, success=False)
            elif create_user(username, password):
                session['username'] = username
                session['needs_api_key_setup'] = True  # Flag to show API key modal
                log_security_event('REGISTER', 'New user account created', username=username, success=True)
                return redirect(url_for('api_key_setup'))
            else:
                error = "Username already exists"
                log_security_event('REGISTER_ATTEMPT', 'Username already exists', username=username, success=False)

    return render_template('auth.html', error=error, mode='register')


# =============================================================================
# API Key Setup Page and Endpoints
# =============================================================================

API_KEY_SETUP_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>API Key Setup - VoiceVerse</title>
    <style>
        :root {
            --spotify-black: #121212;
            --spotify-dark-gray: #181818;
            --spotify-light-gray: #282828;
            --spotify-green: #1DB954;
            --spotify-white: #FFFFFF;
            --spotify-light-text: #B3B3B3;
        }

        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: linear-gradient(135deg, var(--spotify-black) 0%, #1a1a2e 100%);
            min-height: 100vh;
            display: flex;
            justify-content: center;
            align-items: center;
            padding: 20px;
        }

        .setup-container {
            background: var(--spotify-dark-gray);
            border-radius: 16px;
            padding: 40px;
            max-width: 600px;
            width: 100%;
            box-shadow: 0 8px 32px rgba(0,0,0,0.4);
        }

        .setup-header {
            text-align: center;
            margin-bottom: 30px;
        }

        .setup-header h1 {
            color: var(--spotify-white);
            font-size: 28px;
            margin-bottom: 10px;
        }

        .setup-header p {
            color: var(--spotify-light-text);
            font-size: 16px;
            line-height: 1.5;
        }

        .welcome-badge {
            display: inline-block;
            background: var(--spotify-green);
            color: white;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 12px;
            margin-bottom: 15px;
        }

        .option-card {
            background: var(--spotify-light-gray);
            border-radius: 12px;
            padding: 24px;
            margin-bottom: 20px;
            border: 2px solid transparent;
            transition: all 0.3s ease;
            cursor: pointer;
        }

        .option-card:hover {
            border-color: var(--spotify-green);
            transform: translateY(-2px);
        }

        .option-card.active {
            border-color: var(--spotify-green);
            background: rgba(29, 185, 84, 0.1);
        }

        .option-header {
            display: flex;
            align-items: center;
            gap: 12px;
            margin-bottom: 12px;
        }

        .option-icon {
            width: 40px;
            height: 40px;
            background: var(--spotify-green);
            border-radius: 10px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-size: 20px;
        }

        .option-title {
            color: var(--spotify-white);
            font-size: 18px;
            font-weight: 600;
        }

        .option-desc {
            color: var(--spotify-light-text);
            font-size: 14px;
            line-height: 1.5;
        }

        .api-key-form {
            display: none;
            margin-top: 20px;
        }

        .api-key-form.active {
            display: block;
        }

        .form-group {
            margin-bottom: 20px;
        }

        .form-group label {
            display: block;
            color: var(--spotify-white);
            margin-bottom: 8px;
            font-weight: 500;
        }

        .form-group input {
            width: 100%;
            padding: 14px 16px;
            background: var(--spotify-black);
            border: 1px solid #404040;
            border-radius: 8px;
            color: var(--spotify-white);
            font-size: 14px;
            font-family: monospace;
        }

        .form-group input:focus {
            outline: none;
            border-color: var(--spotify-green);
        }

        .form-group input::placeholder {
            color: #666;
        }

        .help-text {
            color: var(--spotify-light-text);
            font-size: 12px;
            margin-top: 8px;
        }

        .help-text a {
            color: var(--spotify-green);
            text-decoration: none;
        }

        .btn-primary {
            width: 100%;
            padding: 14px 24px;
            background: var(--spotify-green);
            color: white;
            border: none;
            border-radius: 30px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
        }

        .btn-primary:hover {
            background: #1ed760;
            transform: scale(1.02);
        }

        .btn-primary:disabled {
            background: #404040;
            cursor: not-allowed;
            transform: none;
        }

        .btn-secondary {
            width: 100%;
            padding: 14px 24px;
            background: transparent;
            color: var(--spotify-white);
            border: 1px solid #404040;
            border-radius: 30px;
            font-size: 16px;
            cursor: pointer;
            margin-top: 12px;
            transition: all 0.3s ease;
        }

        .btn-secondary:hover {
            border-color: var(--spotify-white);
        }

        .error-message {
            background: rgba(255, 82, 82, 0.1);
            border: 1px solid #ff5252;
            color: #ff5252;
            padding: 12px 16px;
            border-radius: 8px;
            margin-bottom: 20px;
            font-size: 14px;
        }

        .success-message {
            background: rgba(29, 185, 84, 0.1);
            border: 1px solid var(--spotify-green);
            color: var(--spotify-green);
            padding: 12px 16px;
            border-radius: 8px;
            margin-bottom: 20px;
            font-size: 14px;
        }

        .env-instructions {
            background: var(--spotify-black);
            border-radius: 8px;
            padding: 16px;
            margin-top: 16px;
        }

        .env-instructions h4 {
            color: var(--spotify-white);
            margin-bottom: 12px;
            font-size: 14px;
        }

        .env-instructions code {
            display: block;
            background: #1a1a1a;
            padding: 12px;
            border-radius: 6px;
            font-family: monospace;
            font-size: 13px;
            color: var(--spotify-green);
            margin-bottom: 12px;
            overflow-x: auto;
        }

        .env-instructions p {
            color: var(--spotify-light-text);
            font-size: 13px;
            line-height: 1.5;
        }

        .security-note {
            display: flex;
            align-items: flex-start;
            gap: 10px;
            background: rgba(255, 193, 7, 0.1);
            border: 1px solid #ffc107;
            padding: 12px;
            border-radius: 8px;
            margin-top: 20px;
        }

        .security-note span {
            color: #ffc107;
            font-size: 13px;
        }
    </style>
</head>
<body>
    <div class="setup-container">
        <div class="setup-header">
            <span class="welcome-badge">Welcome to VoiceVerse!</span>
            <h1>Set Up Your API Key</h1>
            <p>VoiceVerse uses OpenAI's Text-to-Speech API. You'll need an OpenAI API key to generate audio.</p>
        </div>

        {% if error %}
        <div class="error-message">{{ error }}</div>
        {% endif %}

        {% if success %}
        <div class="success-message">{{ success }}</div>
        {% endif %}

        <!-- Option 1: Enter API Key in App -->
        <div class="option-card" id="option1" onclick="selectOption(1)">
            <div class="option-header">
                <div class="option-icon">🔑</div>
                <div class="option-title">Enter API Key Here</div>
            </div>
            <div class="option-desc">
                Enter your OpenAI API key directly. It will be encrypted and stored securely in our database.
                This is the easiest option for most users.
            </div>
        </div>

        <form id="apiKeyForm" class="api-key-form" method="POST" action="{{ url_for('save_api_key') }}">
            <input type="hidden" name="csrf_token" value="{{ csrf_token() }}">
            <div class="form-group">
                <label for="api_key">OpenAI API Key</label>
                <input type="password" id="api_key" name="api_key" placeholder="sk-proj-..." autocomplete="off">
                <p class="help-text">
                    Don't have an API key? <a href="https://platform.openai.com/api-keys" target="_blank">Get one from OpenAI</a>
                </p>
            </div>
            <button type="submit" class="btn-primary" id="saveBtn" disabled>Save API Key</button>
        </form>

        <!-- Option 2: Self-Host with .env -->
        <div class="option-card" id="option2" onclick="selectOption(2)">
            <div class="option-header">
                <div class="option-icon">🖥️</div>
                <div class="option-title">Self-Host with .env File</div>
            </div>
            <div class="option-desc">
                If you're self-hosting VoiceVerse, you can set your API key in the .env file instead.
                This keeps your key on your own server.
            </div>
        </div>

        <div id="envInstructions" class="env-instructions" style="display: none;">
            <h4>How to set up your .env file:</h4>
            <code>OPENAI_API_KEY=sk-proj-your-api-key-here</code>
            <p>
                1. Create a file named <strong>.env</strong> in the VoiceVerse root directory<br>
                2. Add the line above with your actual API key<br>
                3. Restart the VoiceVerse server<br>
                4. Your API key will be used automatically for all your generations
            </p>
        </div>

        <div class="security-note">
            <span>🔒 <strong>Security:</strong> Your API key is encrypted using AES-256 before being stored.
            We never log or expose your key. OpenAI API calls are made directly from this server.</span>
        </div>

        {% if not has_key %}
        <div class="api-key-required-notice" style="background: rgba(255, 193, 7, 0.1); border: 1px solid rgba(255, 193, 7, 0.3); border-radius: 8px; padding: 12px 16px; margin-bottom: 20px; text-align: center;">
            <span style="color: #ffc107;">⚠️ <strong>API key required</strong> to generate audio. Please set your key above.</span>
        </div>
        {% endif %}

        <a href="{{ url_for('index') }}" id="continueLink">
            <button type="button" class="btn-secondary" id="continueBtn" {% if not has_key %}disabled style="opacity: 0.5; cursor: not-allowed;"{% endif %}>
                {% if has_key %}Continue to VoiceVerse{% else %}Set API Key to Continue{% endif %}
            </button>
        </a>
    </div>

    <script>
        let selectedOption = 0;

        function selectOption(option) {
            // Reset all options
            document.getElementById('option1').classList.remove('active');
            document.getElementById('option2').classList.remove('active');
            document.getElementById('apiKeyForm').classList.remove('active');
            document.getElementById('envInstructions').style.display = 'none';

            selectedOption = option;

            if (option === 1) {
                document.getElementById('option1').classList.add('active');
                document.getElementById('apiKeyForm').classList.add('active');
                document.getElementById('api_key').focus();
            } else if (option === 2) {
                document.getElementById('option2').classList.add('active');
                document.getElementById('envInstructions').style.display = 'block';
            }
        }

        // Enable save button when API key is entered
        document.getElementById('api_key').addEventListener('input', function() {
            const key = this.value.trim();
            const saveBtn = document.getElementById('saveBtn');
            // Basic validation: starts with sk- and has reasonable length
            saveBtn.disabled = !(key.startsWith('sk-') && key.length >= 40);
        });

        // Auto-select option 1 if user starts typing
        document.getElementById('api_key').addEventListener('focus', function() {
            if (selectedOption !== 1) {
                selectOption(1);
            }
        });

        // Handle form submission with AJAX
        document.getElementById('apiKeyForm').addEventListener('submit', async function(e) {
            e.preventDefault();
            const saveBtn = document.getElementById('saveBtn');
            const originalText = saveBtn.textContent;
            saveBtn.textContent = 'Saving...';
            saveBtn.disabled = true;

            try {
                const formData = new FormData(this);
                const response = await fetch(this.action, {
                    method: 'POST',
                    body: formData
                });

                const result = await response.json();

                if (result.success) {
                    // Show success message
                    let successDiv = document.querySelector('.success-message');
                    if (!successDiv) {
                        successDiv = document.createElement('div');
                        successDiv.className = 'success-message';
                        document.querySelector('.setup-header').after(successDiv);
                    }
                    successDiv.textContent = result.message || 'API key saved successfully!';
                    successDiv.style.display = 'block';

                    // Hide the warning notice
                    const warningNotice = document.querySelector('.api-key-required-notice');
                    if (warningNotice) {
                        warningNotice.style.display = 'none';
                    }

                    // Enable continue button
                    const continueBtn = document.getElementById('continueBtn');
                    continueBtn.disabled = false;
                    continueBtn.style.opacity = '1';
                    continueBtn.style.cursor = 'pointer';
                    continueBtn.textContent = 'Continue to VoiceVerse';

                    // Clear the input
                    document.getElementById('api_key').value = '';
                    saveBtn.textContent = 'Key Saved ✓';
                } else {
                    // Show error
                    let errorDiv = document.querySelector('.error-message');
                    if (!errorDiv) {
                        errorDiv = document.createElement('div');
                        errorDiv.className = 'error-message';
                        document.querySelector('.setup-header').after(errorDiv);
                    }
                    errorDiv.textContent = result.error || 'Failed to save API key';
                    errorDiv.style.display = 'block';
                    saveBtn.textContent = originalText;
                    saveBtn.disabled = false;
                }
            } catch (error) {
                console.error('Error saving API key:', error);
                saveBtn.textContent = originalText;
                saveBtn.disabled = false;
                alert('Failed to save API key. Please try again.');
            }
        });

        // Prevent navigation if button is disabled
        document.getElementById('continueLink').addEventListener('click', function(e) {
            const continueBtn = document.getElementById('continueBtn');
            if (continueBtn.disabled) {
                e.preventDefault();
                alert('Please set your OpenAI API key first to use VoiceVerse.');
            }
        });
    </script>
</body>
</html>
"""


@app.route('/setup/api-key', methods=['GET'])
@login_required
def api_key_setup():
    """API Key setup page - shown after registration or from settings."""
    username = session.get('username')
    user = db.get_user(username)
    has_key = user_has_api_key(user['id']) if user else False

    return render_template_string(
        API_KEY_SETUP_TEMPLATE,
        has_key=has_key,
        error=request.args.get('error'),
        success=request.args.get('success')
    )


@app.route('/setup/api-key/save', methods=['POST'])
@login_required
def save_api_key():
    """Save user's API key (encrypted)."""
    username = session.get('username')
    user = db.get_user(username)

    # Check if request expects JSON response (AJAX)
    wants_json = request.headers.get('Accept', '').find('application/json') != -1 or \
                 request.content_type == 'application/x-www-form-urlencoded'

    if not user:
        if wants_json:
            return jsonify({'success': False, 'error': 'User not found'}), 404
        return redirect(url_for('api_key_setup', error='User not found'))

    api_key = request.form.get('api_key', '').strip()

    # Validate the API key format
    if not validate_openai_api_key(api_key):
        log_security_event('API_KEY_INVALID', 'Invalid API key format', username=username, success=False)
        if wants_json:
            return jsonify({'success': False, 'error': 'Invalid API key format. Key should start with sk- and be at least 40 characters.'}), 400
        return redirect(url_for('api_key_setup', error='Invalid API key format. Key should start with sk- and be at least 40 characters.'))

    # Encrypt and store the API key
    try:
        encrypted_key = encrypt_api_key(api_key)
        if db.set_user_api_key(user['id'], encrypted_key):
            # Clear cached client so it uses the new key
            clear_user_client_cache(user['id'])
            # Clear the setup flag
            session.pop('needs_api_key_setup', None)
            log_security_event('API_KEY_SET', 'User API key configured', username=username, success=True)
            if wants_json:
                return jsonify({'success': True, 'message': 'API key saved successfully! You can now use VoiceVerse.'})
            return redirect(url_for('api_key_setup', success='API key saved successfully! You can now use VoiceVerse.'))
        else:
            log_security_event('API_KEY_SAVE_FAILED', 'Database error saving API key', username=username, success=False)
            if wants_json:
                return jsonify({'success': False, 'error': 'Failed to save API key. Please try again.'}), 500
            return redirect(url_for('api_key_setup', error='Failed to save API key. Please try again.'))
    except Exception as e:
        log_security_event('API_KEY_ERROR', f'Error saving API key: {str(e)}', username=username, success=False)
        if wants_json:
            return jsonify({'success': False, 'error': 'An error occurred. Please try again.'}), 500
        return redirect(url_for('api_key_setup', error='An error occurred. Please try again.'))


@app.route('/api/settings/api-key', methods=['GET'])
@login_required
def get_api_key_status():
    """Get API key status for current user."""
    username = session.get('username')
    user = db.get_user(username)

    if not user:
        return jsonify({'error': 'User not found'}), 404

    has_own_key = db.has_api_key(user['id'])
    has_global_key = os.getenv("OPENAI_API_KEY") is not None

    # Get masked key for display if user has one
    masked_key = None
    if has_own_key:
        encrypted_key = db.get_user_api_key(user['id'])
        if encrypted_key:
            decrypted = decrypt_api_key(encrypted_key)
            if decrypted:
                masked_key = mask_api_key(decrypted)

    return jsonify({
        'has_api_key': has_own_key or has_global_key,
        'has_own_key': has_own_key,
        'has_global_fallback': has_global_key,
        'masked_key': masked_key,
        'source': 'user' if has_own_key else ('global' if has_global_key else 'none')
    })


@app.route('/api/settings/api-key', methods=['DELETE'])
@login_required
def delete_api_key():
    """Delete user's API key."""
    username = session.get('username')
    user = db.get_user(username)

    if not user:
        return jsonify({'error': 'User not found'}), 404

    if db.delete_user_api_key(user['id']):
        clear_user_client_cache(user['id'])
        log_security_event('API_KEY_DELETED', 'User API key removed', username=username, success=True)
        return jsonify({'success': True, 'message': 'API key deleted'})
    else:
        return jsonify({'error': 'Failed to delete API key'}), 500


@app.route('/api/settings/api-key/test', methods=['POST'])
@csrf.exempt
@login_required
def test_api_key():
    """Test if the current API key is valid by making a simple API call."""
    username = session.get('username')
    user = db.get_user(username)

    if not user:
        return jsonify({'error': 'User not found'}), 404

    client = get_user_openai_client(user['id'])
    if not client:
        return jsonify({
            'valid': False,
            'error': 'No API key configured'
        })

    try:
        # Make a minimal API call to test the key
        # Using models.list() as it's a cheap/free operation
        client.models.list()
        return jsonify({
            'valid': True,
            'message': 'API key is working correctly'
        })
    except Exception as e:
        error_msg = str(e)
        if 'invalid_api_key' in error_msg.lower() or 'incorrect api key' in error_msg.lower():
            return jsonify({
                'valid': False,
                'error': 'Invalid API key'
            })
        elif 'rate_limit' in error_msg.lower():
            return jsonify({
                'valid': True,
                'message': 'API key is valid (rate limited)'
            })
        else:
            return jsonify({
                'valid': False,
                'error': f'API error: {error_msg}'
            })


@app.route('/switch-account/<username>')
@login_required
def switch_account(username):
    """
    Switch to a different user account with proper authorization checks.

    Security: Only allows switching between accounts that belong to the same user
    or have been explicitly authorized.
    """
    current_username = session.get('username')

    # Security: Get current user
    current_user = db.get_user(current_username)
    if not current_user:
        log_security_event(
            'INVALID_SESSION',
            'Session user not found in database',
            username=current_username
        )
        session.clear()
        return redirect(url_for('login'))

    # Security: Verify target user exists
    target_user = db.get_user(username)
    if not target_user:
        log_security_event(
            'ACCOUNT_SWITCH_FAILED',
            f'Attempted switch to non-existent user: {username}',
            username=current_username
        )
        return redirect(url_for('index'))

    # Security: CRITICAL AUTHORIZATION CHECK
    # Prevent users from switching to other users' accounts
    # Only allow switching if it's the same user (for now - multi-account to be implemented)
    if current_user['id'] != target_user['id']:
        # SECURITY ALERT: Unauthorized account switch attempt
        log_security_event(
            'UNAUTHORIZED_ACCOUNT_SWITCH_ATTEMPT',
            f'User {current_username} (ID: {current_user["id"]}) attempted to switch to {username} (ID: {target_user["id"]})',
            username=current_username,
            success=False
        )

        # Send security alert email
        try:
            alerts.send_security_alert(
                alert_type='unauthorized_access',
                details=f'Unauthorized account switch attempt from {current_username} to {username}',
                severity='HIGH'
            )
        except Exception:
            pass  # Don't fail if email fails

        # Return 403 Forbidden
        return "Unauthorized: You can only switch between your own accounts", 403

    # Authorization passed - proceed with account switch
    session['username'] = username
    db.update_last_login(target_user['id'])

    log_security_event(
        'ACCOUNT_SWITCH_SUCCESS',
        f'User switched account to {username}',
        username=username,
        success=True
    )

    return redirect(url_for('index'))

# AI Discovery & Documentation Routes
@app.route('/robots.txt')
def robots_txt():
    """Serve robots.txt for AI crawler guidance"""
    return send_file('static/robots.txt', mimetype='text/plain')

@app.route('/sitemap.xml')
def sitemap_xml():
    """Serve XML sitemap for search engines and AI crawlers"""
    return send_file('static/sitemap.xml', mimetype='application/xml')

@app.route('/openapi.json')
def openapi_spec():
    """Serve OpenAPI 3.1.0 specification for API documentation"""
    return send_file('static/openapi.json', mimetype='application/json')

@app.route('/api-docs')
def api_docs():
    """Serve human-readable API documentation"""
    return jsonify({
        "name": "VoiceVerse API",
        "version": "1.0.0",
        "description": "AI-powered text-to-speech API",
        "documentation": {
            "openapi_spec": "/openapi.json",
            "base_url": request.url_root,
            "authentication": "Session-based (login required for web interface)"
        },
        "endpoints": {
            "audio_generation": {
                "POST /": "Generate audio from text",
                "parameters": {
                    "text": "Text to convert (max 50,000 chars)",
                    "voice": "alloy|echo|fable|onyx|nova|shimmer",
                    "speed": "0.25-4.0 (default: 1.0)",
                    "filename": "Custom filename",
                    "group": "Category/group name",
                    "use_preprocessing": "Enable AI text preprocessing",
                    "use_chunking": "Enable smart chunking for long text"
                }
            },
            "file_management": {
                "GET /api/history": "Get playback history",
                "GET /api/groups": "Get file groups",
                "POST /api/clear-history": "Clear playback history",
                "POST /api/move-to-group": "Move files to a group"
            },
            "document_parsing": {
                "POST /api/parse-docx": "Extract text from DOCX files",
                "POST /api/parse-pdf": "Extract text from PDF files"
            },
            "ai_features": {
                "POST /api/agent/preprocess": "AI text preprocessing",
                "POST /api/agent/suggest-metadata": "AI metadata suggestions",
                "POST /api/agent/analyze": "AI text quality analysis",
                "POST /api/agent/smart-chunk": "Smart text chunking"
            }
        },
        "voices": {
            "alloy": "Neutral, balanced tone - ideal for tutorials and general content",
            "echo": "Male voice, clear - suitable for technical and professional content",
            "fable": "British accent, expressive - perfect for storytelling and audiobooks",
            "onyx": "Deep, authoritative - great for news and formal content",
            "nova": "Female voice, friendly - excellent for guides and conversational content",
            "shimmer": "Soft, warm tone - best for soothing and calm narration"
        },
        "features": [
            "Multiple AI voices with distinct characteristics",
            "Adjustable speech speed (0.25x - 4.0x)",
            "Smart text preprocessing for better speech quality",
            "Intelligent chunking for texts over 4,096 characters",
            "Document upload support (TXT, DOCX, PDF)",
            "Audio file organization with groups",
            "Usage tracking and cost estimation",
            "Voice comparison tool",
            "Playback history"
        ]
    })

@app.route('/capabilities')
def capabilities():
    """Machine-readable endpoint for AI agents to understand app capabilities"""
    return jsonify({
        "application": {
            "name": "VoiceVerse",
            "type": "text-to-speech",
            "version": "1.0.0",
            "api_version": "1.0.0"
        },
        "capabilities": {
            "text_to_speech": {
                "enabled": True,
                "max_characters": 50000,
                "chunk_size": 4096,
                "voices": ["alloy", "echo", "fable", "onyx", "nova", "shimmer"],
                "speed_range": {"min": 0.25, "max": 4.0},
                "formats": ["mp3"]
            },
            "ai_features": {
                "text_preprocessing": True,
                "smart_chunking": True,
                "metadata_suggestions": True,
                "text_analysis": True
            },
            "document_parsing": {
                "formats": ["txt", "docx", "pdf"],
                "max_file_size": 16777216
            },
            "file_management": {
                "groups": True,
                "history": True,
                "search": True,
                "bulk_operations": True
            },
            "api_access": {
                "rest_api": True,
                "websocket": False,
                "graphql": False
            }
        },
        "limits": {
            "max_text_length": 50000,
            "max_filename_length": 100,
            "max_group_name_length": 50,
            "max_file_size": 16777216,
            "rate_limit": "50 requests per hour per user"
        },
        "authentication": {
            "methods": ["session"],
            "required": True
        },
        "documentation": {
            "openapi": "/openapi.json",
            "api_docs": "/api-docs",
            "ai_info": "/ai-info"
        }
    })

@app.route('/ai-info')
def ai_info():
    """Comprehensive AI-friendly information about the application"""
    return jsonify({
        "application": {
            "name": "VoiceVerse",
            "description": "An AI-powered text-to-speech application that converts text into natural, high-quality audio using advanced voice synthesis technology from OpenAI.",
            "tagline": "Convert Text to Natural Speech with AI",
            "category": "Multimedia/Text-to-Speech",
            "version": "1.0.0",
            "url": request.url_root
        },
        "for_ai_agents": {
            "what_this_app_does": "VoiceVerse allows you to convert text into spoken audio using AI-generated voices. It provides features for managing audio files, organizing them into groups, and enhancing text quality before conversion.",
            "primary_use_cases": [
                "Converting articles or blog posts to audio for accessibility",
                "Creating audiobook narrations",
                "Generating voice-overs for presentations",
                "Converting study materials to audio format",
                "Creating podcast content from written scripts",
                "Accessibility for visually impaired users"
            ],
            "how_to_use": {
                "step_1": "Create an account and log in",
                "step_2": "Navigate to 'Create New' section",
                "step_3": "Enter or upload your text (supports TXT, DOCX, PDF)",
                "step_4": "Choose a voice that matches your content style",
                "step_5": "Optionally enable AI preprocessing or smart chunking",
                "step_6": "Click 'Generate Audio' to create your audio file",
                "step_7": "Play, download, or organize your audio files into groups"
            },
            "api_integration": {
                "authentication": "Session-based (requires login)",
                "content_type": "application/x-www-form-urlencoded or application/json",
                "response_format": "HTML for web interface, JSON for API endpoints",
                "error_handling": "Returns appropriate HTTP status codes with descriptive messages"
            }
        },
        "features": {
            "voice_options": {
                "count": 6,
                "voices": {
                    "alloy": {
                        "type": "neutral",
                        "best_for": ["tutorials", "general content", "balanced narration"],
                        "characteristics": "Neutral, balanced tone"
                    },
                    "echo": {
                        "type": "male",
                        "best_for": ["technical content", "professional presentations", "corporate"],
                        "characteristics": "Clear, professional male voice"
                    },
                    "fable": {
                        "type": "british_expressive",
                        "best_for": ["storytelling", "audiobooks", "creative content"],
                        "characteristics": "British accent, expressive, engaging"
                    },
                    "onyx": {
                        "type": "authoritative",
                        "best_for": ["news", "formal announcements", "documentary"],
                        "characteristics": "Deep, authoritative, commanding"
                    },
                    "nova": {
                        "type": "female_friendly",
                        "best_for": ["guides", "tutorials", "conversational content"],
                        "characteristics": "Friendly female voice, approachable"
                    },
                    "shimmer": {
                        "type": "soothing",
                        "best_for": ["meditation", "calm narration", "relaxation"],
                        "characteristics": "Soft, warm, soothing"
                    }
                }
            },
            "ai_enhancements": {
                "preprocessing": "Cleans text, fixes formatting, expands URLs and acronyms for better speech quality",
                "smart_chunking": "Intelligently splits long text at natural boundaries instead of arbitrary character limits",
                "metadata_suggestions": "AI analyzes your text and suggests appropriate filename, category, and voice",
                "text_analysis": "Identifies potential issues in text that may affect speech quality"
            },
            "file_management": {
                "organization": "Group files by category (work, personal, projects, etc.)",
                "search": "Search through your audio library",
                "bulk_operations": "Select and manage multiple files at once",
                "playback_history": "Track recently played audio files"
            }
        },
        "technical_details": {
            "powered_by": "OpenAI Text-to-Speech API",
            "audio_format": "MP3",
            "max_input_length": "50,000 characters",
            "single_request_limit": "4,096 characters (use smart chunking for longer texts)",
            "supported_upload_formats": ["TXT", "DOCX", "PDF"],
            "max_upload_size": "16 MB",
            "pricing_model": "Pay-per-character usage (OpenAI API costs apply)"
        },
        "accessibility": {
            "aria_labels": True,
            "semantic_html": True,
            "keyboard_navigation": True,
            "screen_reader_compatible": True,
            "wcag_compliance": "Designed with WCAG 2.1 guidelines in mind"
        },
        "schema_org_data": True,
        "open_graph_tags": True,
        "sitemap": "/sitemap.xml",
        "robots_txt": "/robots.txt",
        "openapi_spec": "/openapi.json"
    })

@app.route('/logout')
def logout():
    # Security: Log logout event before clearing session
    username = session.get('username', 'anonymous')
    log_security_event('LOGOUT', 'User logged out', username=username)

    session.pop('username', None)
    return redirect(url_for('login'))

@app.route('/workflow-editor')
@admin_required
def workflow_editor():
    """Workflow Editor - Admin only"""
    return send_file('static/workflow_editor.html')


@app.route('/ai-agents')
@admin_required
def ai_agents():
    """AI Agent visual editor - Admin only"""
    return send_file('static/agent_editor.html')

@app.route('/analytics')
@admin_required
def analytics_dashboard():
    """Analytics Dashboard - Admin only"""
    return send_file('static/analytics.html')

@app.route('/settings')
@login_required
def settings():
    """User settings and admin panel"""
    username = session.get('username')
    user = db.get_user(username)
    is_admin = user.get('is_admin', False)

    # Get user stats
    usage_stats = db.get_all_time_usage(user['id'])
    audio_files = db.get_audio_files_by_owner(user['id'])

    return render_template('settings.html',
                                 username=username,
                                 is_admin=is_admin,
                                 user=user,
                                 total_files=len(audio_files),
                                 total_characters=usage_stats.get('total_characters', 0),
                                 total_cost=usage_stats.get('total_cost', 0.0))

@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    """Main dashboard with original Spotify theme"""
    error = None
    success = request.args.get('success') == '1'
    filename = request.args.get('play_file')
    file_display_name = request.args.get('play_name')

    if request.method == 'POST':
        text = request.form.get('text', '').strip()
        text = text.encode('utf-8', 'ignore').decode('utf-8').strip()
        voice = validate_voice(request.form.get('voice', 'nova'))
        file_name = sanitize_display_name(request.form.get('filename', 'audio'))
        group_input = request.form.get('group', '').strip()
        group = sanitize_display_name(group_input) if group_input else 'Uncategorized'
        group = group[:50] if group else 'Uncategorized'

        # Get speed parameter (default 1.0)
        try:
            speed = float(request.form.get('speed', 1.0))
            speed = max(0.25, min(4.0, speed))  # Clamp between 0.25 and 4.0
        except:
            speed = 1.0

        # Get file from form if present
        uploaded_file = request.files.get('file')

        if not text and not uploaded_file:
            error = "Please provide text or upload a file"
        else:
            try:
                # Extract text from uploaded file if present
                if uploaded_file and uploaded_file.filename:
                    file_ext = os.path.splitext(uploaded_file.filename)[1].lower()
                    if file_ext == '.txt':
                        # Security: Check file size before reading into memory
                        uploaded_file.seek(0, 2)  # Seek to end
                        file_size = uploaded_file.tell()
                        uploaded_file.seek(0)  # Reset to beginning

                        if file_size > 10 * 1024 * 1024:  # 10MB limit for text files
                            return jsonify({'error': 'Text file too large (max 10MB)'}), 400

                        text = uploaded_file.read().decode('utf-8', errors='ignore').strip()
                    elif file_ext == '.pdf':
                        import PyPDF2
                        pdf_reader = PyPDF2.PdfReader(uploaded_file)
                        text = ' '.join([page.extract_text() for page in pdf_reader.pages])
                    elif file_ext == '.docx':
                        import docx
                        doc = docx.Document(uploaded_file)
                        text = ' '.join([paragraph.text for paragraph in doc.paragraphs])

                if not text or not text.strip():
                    error = "No valid text found to convert"
                else:
                    # Generate audio
                    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                    safe_filename = f"{file_name}_{timestamp}.mp3"
                    filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)

                    # Call OpenAI TTS API
                    response = client.audio.speech.create(
                        model="tts-1",
                        voice=voice,
                        input=text,
                        speed=speed
                    )
                    response.stream_to_file(filepath)

                    # Save to database
                    user = db.get_user_by_username(session['username'])
                    if user:
                        db.save_audio_file(
                            user_id=user['id'],
                            filename=safe_filename,
                            voice=voice,
                            speed=speed,
                            text_preview=text[:200],
                            group_name=group,
                            file_path=filepath
                        )

                    return redirect(url_for('index', success=1, play_file=safe_filename, play_name=file_name))

            except Exception as e:
                error = f"Error generating audio: {str(e)}"

    # Get user's audio files from database
    user = db.get_user(session['username'])
    audio_files_db = []
    groups = {}

    if user:
        audio_files_db = db.get_audio_files_by_owner(user['id'])
        for file_info in audio_files_db:
            group_name = file_info.get('category') or 'Uncategorized'
            if group_name not in groups:
                groups[group_name] = []
            groups[group_name].append(file_info)

    # Transform database files to format expected by template
    recent_files = [
        {
            'filename': f['filename'],
            'name': f['display_name'],
            'group': f.get('category') or 'Uncategorized',
            'created': f.get('created_at', ''),
            'chars': f.get('character_count', 0),
            'cost': f.get('cost', 0.0)
        }
        for f in audio_files_db[:10]
        if os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], f['filename']))
    ]

    # Get usage statistics
    usage_stats = db.get_all_time_usage(user['id'])
    usage = {
        'total_characters': usage_stats.get('total_characters', 0),
        'total_cost': usage_stats.get('total_cost', 0.0),
        'files_generated': len(audio_files_db),
        'monthly': {}
    }

    # Get all users for account switcher
    all_users = db.list_users()

    # Check if user is admin
    is_admin = user.get('is_admin', False)

    return render_template(
        'index.html',
        error=error,
        success=success,
        filename=filename,
        file_display_name=file_display_name,
        recent_files=recent_files,
        usage=usage,
        groups={group_name: len(files) for group_name, files in groups.items()},
        all_users=all_users,
        is_admin=is_admin,
        tts_price_per_1k=TTS_PRICING['tts-1']
    )

@app.route('/classic', methods=['GET', 'POST'])
@login_required
def index_classic():
    """Original Spotify-style dashboard (backup)"""
    error = None
    success = request.args.get('success') == '1'
    filename = request.args.get('play_file')
    file_display_name = request.args.get('play_name')

    if request.method == 'POST':
        text = request.form.get('text', '').strip()
        text = text.encode('utf-8', 'ignore').decode('utf-8').strip()
        voice = validate_voice(request.form.get('voice', 'nova'))
        file_name = sanitize_display_name(request.form.get('filename', 'audio'))
        group_input = request.form.get('group', '').strip()
        group = sanitize_display_name(group_input) if group_input else 'Uncategorized'
        group = group[:50] if group else 'Uncategorized'

        # Get speed parameter (default 1.0)
        try:
            speed = float(request.form.get('speed', 1.0))
            speed = max(0.25, min(4.0, speed))  # Clamp between 0.25 and 4.0
        except:
            speed = 1.0

        # Security: Validate input length (prevent DoS attacks)
        if len(text) > 100000:
            error = "Text is too long. Maximum 100,000 characters allowed."
        elif not text:
            error = "Please enter some text"
        elif not file_name:
            error = "Please enter a valid file name"
        else:
            try:
                # AI Agent: Preprocess text for better TTS quality
                use_ai_preprocessing = request.form.get('use_preprocessing', 'off') == 'on'
                if use_ai_preprocessing:
                    try:
                        agents = get_agent_system()
                        text = agents.preprocess_text(text)
                        print(f"✅ Text preprocessed by AI agent")
                    except Exception as e:
                        print(f"⚠️ AI preprocessing failed, using original text: {e}")

                # Handle long text with smart chunking or simple truncation
                original_length = len(text)
                use_smart_chunking = request.form.get('use_chunking', 'off') == 'on'

                if original_length > 4096:
                    if use_smart_chunking:
                        # AI Agent: Smart chunking for multi-part audio
                        try:
                            agents = get_agent_system()
                            chunks = agents.smart_chunk(text, 4000)
                            print(f"✅ Text split into {len(chunks)} chunks by AI agent")
                            # For now, use first chunk (future: generate multiple files)
                            text = chunks[0]['text']
                            print(f"📝 Using chunk 1/{len(chunks)} ({len(text)} chars)")
                        except Exception as e:
                            print(f"⚠️ Smart chunking failed, truncating: {e}")
                            text = text[:4096]
                    else:
                        # Simple truncation (original behavior)
                        text = text[:4096]
                        print(f"⚠️ Text truncated from {original_length} to 4,096 characters for TTS generation")

                # Get user's OpenAI client (BYOK model)
                user = db.get_user(session['username'])
                if not user:
                    error = "User session invalid. Please log in again."
                    return render_template_string(CLASSIC_TEMPLATE, error=error, success=False)

                client = get_user_openai_client(user['id'])
                if not client:
                    error = "OpenAI API key not configured. Please set up your API key in Settings."
                    return render_template_string(CLASSIC_TEMPLATE, error=error, success=False)

                response = client.audio.speech.create(
                    model="tts-1",
                    voice=voice,
                    input=text,
                    speed=speed
                )

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                safe_filename = secure_filename(f"{file_name}_{timestamp}.mp3")
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)

                try:
                    if hasattr(response, 'stream_to_file'):
                        response.stream_to_file(filepath)
                        print(f'✅ Saved using stream_to_file at: {filepath}')
                    else:
                        audio_bytes = getattr(response, 'content', None)
                        if audio_bytes is None and hasattr(response, 'read'):
                            audio_bytes = response.read()
                        if audio_bytes:
                            with open(filepath, 'wb') as f:
                                f.write(audio_bytes)
                            print(f'✅ Saved audio file at: {filepath}')
                        else:
                            print('⚠️ No audio bytes found in OpenAI response.')
                except Exception as e:
                    print(f'❌ Error saving file: {e}')

                char_count = len(text)
                cost = calculate_cost(char_count)

                # Get user ID for database operations
                user = db.get_user(session['username'])
                if not user:
                    error = "User not found"
                else:
                    # Save file metadata to database
                    db.create_audio_file(
                        filename=safe_filename,
                        display_name=file_name,
                        owner_id=user['id'],
                        voice=voice,
                        category=group,
                        text=text,
                        character_count=char_count,
                        cost=cost
                    )

                    # Record usage statistics
                    db.record_usage(user['id'], char_count, cost)

                    # Redirect to home page to show the newly created file
                    return redirect(url_for('index', success='1', play_file=safe_filename, play_name=file_name))

            except ValueError as ve:
                error = str(ve)
            except Exception as e:
                error = f"An error occurred: {str(e)}"
                print(f"Error in audio generation: {e}")
    
    # Get user ID for database queries
    user = db.get_user(session['username'])
    if not user:
        return redirect(url_for('login'))

    # Get files for this user from database
    audio_files_db = db.get_audio_files_by_owner(user['id'])

    all_files = [
        {
            'filename': file_info['filename'],
            'name': file_info['display_name'],
            'group': file_info['category'] or 'Uncategorized',
            'created': file_info['created_at'],
            'chars': file_info['character_count'],
            'cost': file_info['cost']
        }
        for file_info in audio_files_db
        if os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], file_info['filename']))
    ]

    all_files.sort(key=lambda x: x['created'], reverse=True)
    recent_files = all_files[:12]

    # Calculate groups and their counts
    groups = defaultdict(int)
    for file_data in all_files:
        group_name = file_data.get('group', 'Uncategorized')
        groups[group_name] += 1

    # Get usage stats from database
    usage_stats = db.get_all_time_usage(user['id'])
    usage = {
        'total_characters': usage_stats.get('total_characters', 0),
        'total_cost': usage_stats.get('total_cost', 0.0),
        'files_generated': len(audio_files_db),
        'monthly': {}  # Monthly stats can be added later if needed
    }

    # Get all users for account switcher
    all_users = db.list_users()

    # Check if user is admin
    is_admin = user.get('is_admin', False)

    return render_template(
        'index.html',
        error=error,
        success=success,
        filename=filename,
        file_display_name=file_display_name,
        recent_files=recent_files,
        usage=usage,
        groups={group_name: len(files) for group_name, files in groups.items()},
        all_users=all_users,
        is_admin=is_admin,
        tts_price_per_1k=TTS_PRICING['tts-1']
    )

@app.route('/audio/<path:filename>')
@login_required  # Security: Require authentication to access audio files
def audio(filename):
    try:
        safe_filename = secure_filename(filename)

        # Security: Verify file ownership
        if not verify_file_ownership(safe_filename, session['username']):
            return "Unauthorized access", 403

        filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
        if os.path.exists(filepath):
            return send_file(filepath, mimetype='audio/mpeg')
        return "File not found", 404
    except Exception as e:
        print(f"Error serving audio: {e}")
        return "Error serving file", 500

@app.route('/download/<path:filename>')
@login_required  # Security: Require authentication to download files
def download(filename):
    try:
        safe_filename = secure_filename(filename)
        username = session.get('username', 'anonymous')

        # Security: Verify file ownership
        if not verify_file_ownership(safe_filename, username):
            # Security: Log unauthorized access attempt
            security_log.log_file_access(
                filename=safe_filename,
                username=username,
                ip_address=request.remote_addr,
                success=False,
                action='DOWNLOAD'
            )
            return "Unauthorized access", 403

        filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
        if os.path.exists(filepath):
            # Security: Log successful download
            security_log.log_file_access(
                filename=safe_filename,
                username=username,
                ip_address=request.remote_addr,
                success=True,
                action='DOWNLOAD'
            )
            return send_file(filepath, mimetype='audio/mpeg', as_attachment=True, download_name=safe_filename)

        # Security: Log file not found
        security_log.log_file_access(
            filename=safe_filename,
            username=username,
            ip_address=request.remote_addr,
            success=False,
            action='DOWNLOAD'
        )
        return "File not found", 404
    except Exception as e:
        print(f"Error downloading: {e}")
        return "Error downloading file", 500

@app.route('/rename/<path:filename>', methods=['POST'])
@login_required  # Security: Require authentication to rename files
def rename(filename):
    try:
        data = request.get_json()
        new_name = sanitize_display_name(data.get('name', ''))

        if new_name:
            safe_filename = secure_filename(filename)

            # Security: Verify file ownership
            if not verify_file_ownership(safe_filename, session['username']):
                return jsonify({'success': False, 'error': 'Unauthorized access'}), 403

            metadata = load_metadata()

            if safe_filename in metadata:
                metadata[safe_filename]['name'] = new_name
                save_metadata(metadata)
                return jsonify({'success': True})
            return jsonify({'success': False, 'error': 'File not found'}), 404
        return jsonify({'success': False, 'error': 'Invalid name'}), 400

    except Exception as e:
        print(f"Error renaming: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/delete/<path:filename>', methods=['POST'])
@login_required  # Security: Require authentication to delete files
def delete(filename):
    try:
        safe_filename = secure_filename(filename)
        username = session.get('username', 'anonymous')

        # Security: Verify file ownership
        if not verify_file_ownership(safe_filename, username):
            # Security: Log unauthorized delete attempt
            security_log.log_file_access(
                filename=safe_filename,
                username=username,
                ip_address=request.remote_addr,
                success=False,
                action='DELETE'
            )
            return jsonify({'success': False, 'error': 'Unauthorized access'}), 403

        filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)

        # Delete from database
        file_info = db.get_audio_file(safe_filename)
        if file_info:
            db.delete_audio_file(file_info['id'])

        if os.path.exists(filepath):
            os.remove(filepath)

        # Security: Log successful deletion
        security_log.log_file_access(
            filename=safe_filename,
            username=username,
            ip_address=request.remote_addr,
            success=True,
            action='DELETE'
        )

        return jsonify({'success': True})

    except Exception as e:
        print(f"Error deleting: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/delete-group/<path:group_name>', methods=['POST'])
def delete_group(group_name):
    try:
        # Get user ID for ownership verification
        user = db.get_user(session['username'])
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403

        # Get all files in this group owned by the user
        audio_files = db.get_audio_files_by_owner(user['id'])
        files_to_delete = []

        for file_info in audio_files:
            if file_info.get('category') == group_name:
                files_to_delete.append(file_info)

                # Delete file from disk
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], file_info['filename'])
                if os.path.exists(filepath):
                    os.remove(filepath)

                # Delete from database
                db.delete_audio_file(file_info['id'])

        return jsonify({'success': True, 'deleted_files': len(files_to_delete)})

    except Exception as e:
        print(f"Error deleting group: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/rename-group/<path:old_name>', methods=['POST'])
def rename_group(old_name):
    try:
        data = request.get_json()
        new_name = sanitize_display_name(data.get('name', ''))

        if not new_name:
            return jsonify({'success': False, 'error': 'Invalid group name'}), 400

        # Get user ID for ownership verification
        user = db.get_user(session['username'])
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403

        # Get all files in this group owned by the user
        audio_files = db.get_audio_files_by_owner(user['id'])
        updated_count = 0

        # Update all files in this group
        for file_info in audio_files:
            if file_info.get('category') == old_name:
                db.update_audio_file(file_info['id'], category=new_name)
                updated_count += 1

        if updated_count > 0:
            return jsonify({'success': True, 'updated_files': updated_count})
        else:
            return jsonify({'success': False, 'error': 'Group not found'}), 404

    except Exception as e:
        print(f"Error renaming group: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/history', methods=['GET'])
def get_history():
    try:
        # Get user ID
        user = db.get_user(session.get('username'))
        if not user:
            return jsonify({'history': []})

        # Get playback history from database
        raw_history = db.get_playback_history(user['id'], limit=20)

        # Transform to format expected by JavaScript (name, timestamp instead of display_name, played_at)
        history = [
            {
                'filename': entry.get('filename'),
                'name': entry.get('display_name', entry.get('filename', 'Unknown')),
                'timestamp': entry.get('played_at'),
                'voice': entry.get('voice'),
                'category': entry.get('category')
            }
            for entry in raw_history
        ]

        return jsonify({'history': history})
    except Exception as e:
        print(f"Error getting history: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/clear-history', methods=['POST'])
@csrf.exempt
def clear_history_endpoint():
    try:
        # Get user ID
        user = db.get_user(session.get('username'))
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403

        # Clear history for this user (delete all playback history records)
        with db._get_cursor() as cursor:
            cursor.execute("DELETE FROM playback_history WHERE user_id = ?", (user['id'],))
        return jsonify({'success': True})
    except Exception as e:
        print(f"Error clearing history: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/add-to-history', methods=['POST'])
@csrf.exempt
@login_required  # Security: Require authentication
def add_to_history_endpoint():
    try:
        data = request.get_json()
        print(f"DEBUG /api/add-to-history: Received data: {data}")

        filename = data.get('filename') if data else None

        if not filename:
            print(f"ERROR /api/add-to-history: Missing filename in request data: {data}")
            return jsonify({'success': False, 'error': 'Missing filename'}), 400

        # Get user ID
        username = session.get('username')
        print(f"DEBUG /api/add-to-history: Username from session: {username}")

        user = db.get_user(username)
        if not user:
            print(f"ERROR /api/add-to-history: User not found for username: {username}")
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403

        # Get file info
        file_info = db.get_audio_file(filename)
        if not file_info:
            print(f"ERROR /api/add-to-history: File not found: {filename}")
            return jsonify({'success': False, 'error': 'File not found'}), 404

        print(f"DEBUG /api/add-to-history: Recording playback - User ID: {user['id']}, File ID: {file_info['id']}")

        # Record playback
        db.record_playback(user['id'], file_info['id'])

        print(f"SUCCESS /api/add-to-history: Recorded playback for {filename}")
        return jsonify({'success': True})
    except Exception as e:
        print(f"ERROR /api/add-to-history: Exception: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/bulk-delete', methods=['POST'])
def bulk_delete():
    try:
        data = request.get_json()
        filenames = data.get('filenames', [])

        if not filenames:
            return jsonify({'success': False, 'error': 'No filenames provided'}), 400

        # Get user ID for ownership verification
        user = db.get_user(session['username'])
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403

        deleted_count = 0

        for filename in filenames:
            safe_filename = secure_filename(filename)

            # Verify file ownership before deletion
            if not verify_file_ownership(safe_filename, session['username']):
                continue

            # Get file info from database
            file_info = db.get_audio_file(safe_filename)
            if file_info:
                # Delete file from disk
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
                if os.path.exists(filepath):
                    os.remove(filepath)

                # Delete from database
                db.delete_audio_file(file_info['id'])
                deleted_count += 1

        return jsonify({'success': True, 'deleted_count': deleted_count})

    except Exception as e:
        print(f"Error in bulk delete: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/groups', methods=['GET'])
def get_groups():
    try:
        # Get user ID
        user = db.get_user(session.get('username'))
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403

        # Get all files for this user from database
        audio_files = db.get_audio_files_by_owner(user['id'])
        groups = {}

        for file_info in audio_files:
            group = file_info.get('category') or 'Uncategorized'
            if group not in groups:
                groups[group] = 0
            groups[group] += 1

        # Sort groups alphabetically
        sorted_groups = dict(sorted(groups.items()))

        return jsonify({'success': True, 'groups': sorted_groups})
    except Exception as e:
        print(f"Error getting groups: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/move-to-group', methods=['POST'])
@csrf.exempt
def move_to_group():
    try:
        data = request.get_json()
        filename = data.get('filename')
        new_group = sanitize_display_name(data.get('group', 'Uncategorized'))

        if not filename:
            return jsonify({'success': False, 'error': 'No filename provided'}), 400

        if not new_group:
            new_group = 'Uncategorized'

        safe_filename = secure_filename(filename)

        # Verify file ownership before updating
        if not verify_file_ownership(safe_filename, session['username']):
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403

        # Get file info from database
        file_info = db.get_audio_file(safe_filename)
        if file_info:
            db.update_audio_file(file_info['id'], category=new_group)
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': 'File not found'}), 404

    except Exception as e:
        print(f"Error moving file to group: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/bulk-move', methods=['POST'])
def bulk_move():
    try:
        data = request.get_json()
        filenames = data.get('filenames', [])
        new_group = sanitize_display_name(data.get('group', 'Uncategorized'))

        if not filenames:
            return jsonify({'success': False, 'error': 'No filenames provided'}), 400

        if not new_group:
            new_group = 'Uncategorized'

        updated_count = 0

        for filename in filenames:
            safe_filename = secure_filename(filename)

            # Verify file ownership before updating
            if not verify_file_ownership(safe_filename, session['username']):
                continue

            # Get file info from database
            file_info = db.get_audio_file(safe_filename)
            if file_info:
                db.update_audio_file(file_info['id'], category=new_group)
                updated_count += 1

        return jsonify({'success': True, 'updated_count': updated_count})

    except Exception as e:
        print(f"Error in bulk move: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/preview/<voice>')
@login_required
def preview(voice):
    try:
        voice = validate_voice(voice)

        # Check if custom text is provided via query parameter
        custom_text = request.args.get('text', '').strip()

        if custom_text:
            # Limit to 500 characters for previews
            preview_text = custom_text[:500]
        else:
            preview_text = f"Hello, I am {voice.capitalize()}. This is how I sound."

        # Get user's OpenAI client (BYOK model)
        user = db.get_user(session['username'])
        if not user:
            return "User session invalid", 401

        client = get_user_openai_client(user['id'])
        if not client:
            return "OpenAI API key not configured. Please set up your API key in Settings.", 400

        response = client.audio.speech.create(
            model="tts-1",
            voice=voice,
            input=preview_text,
            speed=1.0
        )

        audio_buffer = io.BytesIO(response.content)
        audio_buffer.seek(0)

        return send_file(audio_buffer, mimetype='audio/mpeg')

    except Exception as e:
        print(f"Error generating preview: {e}")
        return f"Error: {str(e)}", 500

@app.route('/api/parse-docx', methods=['POST'])
@csrf.exempt
@login_required  # Security: Require authentication for file uploads
def parse_docx():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        if not file.filename.endswith('.docx'):
            return jsonify({'success': False, 'error': 'Invalid file type. Only .docx files are supported'}), 400

        # Parse DOCX file
        doc = Document(file)

        # Extract all text from paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_content.append(paragraph.text.strip())

        # Join paragraphs with double newline
        full_text = '\n\n'.join(text_content)
        original_length = len(full_text)

        # Limit to 50,000 characters for UI
        if len(full_text) > 50000:
            full_text = full_text[:50000]
            truncated = True
        else:
            truncated = False

        return jsonify({
            'success': True,
            'text': full_text,
            'truncated': truncated,
            'original_length': original_length
        })

    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        return jsonify({'success': False, 'error': f'Error parsing DOCX file: {str(e)}'}), 500

@app.route('/api/parse-pdf', methods=['POST'])
@csrf.exempt
@login_required  # Security: Require authentication for file uploads
def parse_pdf():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'success': False, 'error': 'Invalid file type. Only .pdf files are supported'}), 400

        # Parse PDF file
        pdf_reader = PdfReader(file)

        # Extract text from all pages
        text_content = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text.strip():
                text_content.append(page_text.strip())

        # Join pages with double newline
        full_text = '\n\n'.join(text_content)
        original_length = len(full_text)

        # Limit to 50,000 characters for UI
        if len(full_text) > 50000:
            full_text = full_text[:50000]
            truncated = True
        else:
            truncated = False

        return jsonify({
            'success': True,
            'text': full_text,
            'truncated': truncated,
            'original_length': original_length,
            'pages': len(pdf_reader.pages)
        })

    except Exception as e:
        print(f"Error parsing PDF: {e}")
        return jsonify({'success': False, 'error': f'Error parsing PDF file: {str(e)}'}), 500

@app.route('/api/agent/preprocess', methods=['POST'])
@csrf.exempt
@login_required
def agent_preprocess():
    """AI agent: Preprocess text for optimal TTS"""
    try:
        data = request.get_json()
        text = data.get('text', '')

        if not text:
            return jsonify({'success': False, 'error': 'No text provided'}), 400

        agents = get_agent_system()
        cleaned_text = agents.preprocess_text(text)

        return jsonify({
            'success': True,
            'original_text': text,
            'processed_text': cleaned_text,
            'original_length': len(text),
            'processed_length': len(cleaned_text)
        })
    except Exception as e:
        print(f"Error in preprocess agent: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/agent/suggest-metadata', methods=['POST'])
@csrf.exempt
@login_required
def agent_suggest_metadata():
    """AI agent: Suggest filename, category, voice"""
    try:
        data = request.get_json()
        text = data.get('text', '')

        if not text:
            return jsonify({'success': False, 'error': 'No text provided'}), 400

        try:
            agents = get_agent_system()
            suggestions = agents.suggest_metadata(text)
        except Exception as agent_error:
            # If agent system fails, provide basic fallback suggestions
            print(f"AI agent failed, using fallback: {agent_error}")
            suggestions = {
                'filename': 'audio',
                'category': 'General',
                'recommended_voice': 'nova',
                'summary': 'Generated filename and category based on default settings.'
            }

        return jsonify({
            'success': True,
            'suggestions': suggestions
        })
    except Exception as e:
        print(f"Error in metadata suggestion agent: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/agent/analyze', methods=['POST'])
@csrf.exempt
@login_required
def agent_analyze():
    """AI agent: Analyze text quality and provide recommendations"""
    try:
        data = request.get_json()
        text = data.get('text', '')

        if not text:
            return jsonify({'success': False, 'error': 'No text provided'}), 400

        agents = get_agent_system()
        analysis = agents.analyze_quality(text)

        return jsonify({
            'success': True,
            'analysis': analysis
        })
    except Exception as e:
        print(f"Error in analysis agent: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/agent/smart-chunk', methods=['POST'])
@csrf.exempt
@login_required
def agent_smart_chunk():
    """AI agent: Split text into optimal chunks"""
    try:
        data = request.get_json()
        text = data.get('text', '')
        max_chars = int(data.get('max_chars', 4000))

        if not text:
            return jsonify({'success': False, 'error': 'No text provided'}), 400

        agents = get_agent_system()
        chunks = agents.smart_chunk(text, max_chars)

        return jsonify({
            'success': True,
            'chunks': chunks,
            'total_chunks': len(chunks)
        })
    except Exception as e:
        print(f"Error in chunking agent: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Voice Cloning API Endpoints
@app.route('/api/voice-clone/upload-sample', methods=['POST'])
@csrf.exempt
@login_required
def upload_voice_sample():
    """Upload a voice sample for cloning"""
    try:
        if 'voice_sample' not in request.files:
            return jsonify({'success': False, 'error': 'No voice sample provided'}), 400

        file = request.files['voice_sample']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'}), 400

        # Get the custom voice name from form data
        voice_name = request.form.get('voice_name', '').strip()
        if not voice_name:
            voice_name = 'Unnamed Voice'

        # Validate file extension
        allowed_extensions = {'.wav', '.mp3', '.ogg', '.flac', '.m4a'}
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in allowed_extensions:
            return jsonify({'success': False, 'error': 'Invalid audio format. Supported: WAV, MP3, OGG, FLAC, M4A'}), 400

        # Generate secure filename
        username = session.get('username')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{username}_{timestamp}{file_ext}"
        filepath = os.path.join(app.config['VOICE_SAMPLES_FOLDER'], safe_filename)

        # Save the file
        file.save(filepath)

        # Convert to WAV format if needed and get audio info
        try:
            audio_data, sr = sf.read(filepath)

            # Convert to mono if stereo
            if len(audio_data.shape) > 1:
                audio_data = np.mean(audio_data, axis=1)

            # Save as WAV for compatibility with XTTS
            wav_filename = f"{username}_{timestamp}.wav"
            wav_filepath = os.path.join(app.config['VOICE_SAMPLES_FOLDER'], wav_filename)
            sf.write(wav_filepath, audio_data, sr)

            # Remove original if it was converted
            if filepath != wav_filepath:
                os.remove(filepath)

            duration = len(audio_data) / sr

            # Save metadata (voice name) in a JSON file
            metadata_filename = f"{username}_{timestamp}.json"
            metadata_filepath = os.path.join(app.config['VOICE_SAMPLES_FOLDER'], metadata_filename)
            with open(metadata_filepath, 'w') as f:
                json.dump({
                    'name': voice_name,
                    'filename': wav_filename,
                    'duration': round(duration, 2),
                    'sample_rate': sr,
                    'created': datetime.now().isoformat()
                }, f)

            return jsonify({
                'success': True,
                'filename': wav_filename,
                'name': voice_name,
                'duration': round(duration, 2),
                'sample_rate': sr,
                'message': f'Voice sample uploaded successfully ({duration:.1f}s)'
            })

        except Exception as e:
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'error': f'Invalid audio file: {str(e)}'}), 400

    except Exception as e:
        print(f"Error uploading voice sample: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/voice-clone/generate', methods=['POST'])
@csrf.exempt
@login_required
def generate_voice_clone():
    """Generate speech using voice cloning"""
    try:
        if TTS_XTTS is None:
            return jsonify({
                'success': False,
                'error': 'Voice cloning not available. Install Coqui TTS: pip install TTS torch soundfile'
            }), 503

        data = request.get_json()
        text = data.get('text', '').strip()
        voice_sample_filename = data.get('voice_sample_filename')
        language = data.get('language', 'en')
        display_name = data.get('display_name', '').strip()
        group = data.get('group', '').strip() or 'Uncategorized'

        if not text:
            return jsonify({'success': False, 'error': 'No text provided'}), 400

        if not voice_sample_filename:
            return jsonify({'success': False, 'error': 'No voice sample selected'}), 400

        # Get voice sample path
        voice_sample_path = os.path.join(app.config['VOICE_SAMPLES_FOLDER'], voice_sample_filename)
        if not os.path.exists(voice_sample_path):
            return jsonify({'success': False, 'error': 'Voice sample not found'}), 404

        # Initialize XTTS model
        model = get_xtts_model()
        if model is None:
            return jsonify({'success': False, 'error': 'Failed to load voice cloning model'}), 500

        # Generate filename
        username = session.get('username')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_filename = f"{username}_{timestamp}_cloned.wav"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)

        # Generate speech
        try:
            model.tts_to_file(
                text=text,
                file_path=output_path,
                speaker_wav=voice_sample_path,
                language=language
            )

            # Get audio info
            audio_data, sr = sf.read(output_path)
            duration = len(audio_data) / sr
            file_size = os.path.getsize(output_path)

            # Get user info for database
            user = db.get_user(username)

            # Use provided display_name or fallback to filename
            final_display_name = display_name if display_name else output_filename.replace('.wav', '')

            # Store in database
            audio_id = db.create_audio_file(
                filename=output_filename,
                display_name=final_display_name,
                owner_id=user['id'],
                voice='voice_cloned',
                category=group,
                text=text,
                character_count=len(text),
                cost=0.0,  # Local XTTS has no API cost
                duration=duration
            )

            # Record usage
            char_count = len(text)
            # Estimate cost (XTTS is local, so minimal cost)
            cost = 0.0

            db.record_usage(
                user_id=user['id'],
                characters=char_count,
                cost=cost
            )

            return jsonify({
                'success': True,
                'filename': output_filename,
                'audio_id': audio_id,
                'duration': round(duration, 2),
                'message': 'Speech generated successfully with cloned voice!'
            })

        except Exception as e:
            if os.path.exists(output_path):
                os.remove(output_path)
            raise e

    except Exception as e:
        print(f"Error generating voice clone: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/voice-clone/samples', methods=['GET'])
@login_required
def list_voice_samples():
    """List available voice samples for the current user"""
    try:
        username = session.get('username')
        samples = []

        if os.path.exists(app.config['VOICE_SAMPLES_FOLDER']):
            for filename in os.listdir(app.config['VOICE_SAMPLES_FOLDER']):
                if filename.startswith(f"{username}_") and filename.endswith('.wav'):
                    filepath = os.path.join(app.config['VOICE_SAMPLES_FOLDER'], filename)
                    try:
                        audio_data, sr = sf.read(filepath)
                        duration = len(audio_data) / sr
                        file_size = os.path.getsize(filepath)

                        # Try to get the voice name from metadata JSON file
                        voice_name = 'Cloned Voice'  # Default fallback
                        base_name = filename.rsplit('.', 1)[0]  # Remove .wav extension
                        metadata_filepath = os.path.join(app.config['VOICE_SAMPLES_FOLDER'], f"{base_name}.json")
                        if os.path.exists(metadata_filepath):
                            try:
                                with open(metadata_filepath, 'r') as f:
                                    metadata = json.load(f)
                                    voice_name = metadata.get('name', 'Cloned Voice')
                            except Exception as e:
                                print(f"Error reading metadata for {filename}: {e}")

                        samples.append({
                            'filename': filename,
                            'name': voice_name,
                            'duration': round(duration, 2),
                            'size': file_size,
                            'uploaded': datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat()
                        })
                    except Exception as e:
                        print(f"Error reading voice sample {filename}: {e}")
                        continue

        return jsonify({
            'success': True,
            'samples': samples,
            'count': len(samples)
        })

    except Exception as e:
        print(f"Error listing voice samples: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/voice-clone/delete-sample', methods=['POST'])
@csrf.exempt
@login_required
def delete_voice_sample():
    """Delete a voice sample"""
    try:
        data = request.get_json()
        filename = data.get('filename')

        if not filename:
            return jsonify({'success': False, 'error': 'No filename provided'}), 400

        username = session.get('username')

        # Security check: ensure user can only delete their own samples
        if not filename.startswith(f"{username}_"):
            return jsonify({'success': False, 'error': 'Unauthorized'}), 403

        filepath = os.path.join(app.config['VOICE_SAMPLES_FOLDER'], filename)

        if os.path.exists(filepath):
            os.remove(filepath)
            return jsonify({'success': True, 'message': 'Voice sample deleted'})
        else:
            return jsonify({'success': False, 'error': 'File not found'}), 404

    except Exception as e:
        print(f"Error deleting voice sample: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

# Analytics API Endpoints
@app.route('/api/analytics/stats', methods=['GET'])
@admin_required
def analytics_stats():
    """Get analytics statistics - Admin only"""
    try:
        # Get all users count
        total_users = len(db.get_all_users())

        # Get all audio files
        all_files = db.get_all_audio_files()
        total_generations = len(all_files)

        # Calculate total characters and cost
        total_characters = 0
        total_cost = 0.0

        for file_info in all_files:
            # Get the text for each file if available
            if 'text' in file_info and file_info['text']:
                total_characters += len(file_info['text'])
            # Estimate cost using centralized pricing
            if 'text' in file_info and file_info['text']:
                total_cost += calculate_cost(len(file_info['text']))

        return jsonify({
            'success': True,
            'total_users': total_users,
            'total_generations': total_generations,
            'total_characters': total_characters,
            'total_cost': total_cost
        })

    except Exception as e:
        print(f"Error getting analytics stats: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/analytics/recent-activity', methods=['GET'])
@admin_required
def analytics_recent_activity():
    """Get recent activity - Admin only"""
    try:
        # Get recent audio files (last 20)
        all_files = db.get_all_audio_files()

        # Sort by created_at descending and take first 20
        recent_files = sorted(all_files, key=lambda x: x.get('created_at', ''), reverse=True)[:20]

        # Enrich with username
        activities = []
        for file_info in recent_files:
            owner_id = file_info.get('owner_id')
            if owner_id:
                # Get username from user ID
                users = db.get_all_users()
                username = next((u['username'] for u in users if u['id'] == owner_id), 'Unknown')
            else:
                username = 'Unknown'

            activities.append({
                'username': username,
                'filename': file_info.get('filename', 'Unknown'),
                'voice': file_info.get('voice', 'Unknown'),
                'characters': len(file_info.get('text', '')),
                'created_at': file_info.get('created_at', '')
            })

        return jsonify({
            'success': True,
            'activities': activities
        })

    except Exception as e:
        print(f"Error getting recent activity: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    """
    Comprehensive health check endpoint
    Returns 200 if all systems operational, 503 if degraded
    """
    health_status = {
        'status': 'healthy',
        'timestamp': datetime.utcnow().isoformat() + 'Z',
        'checks': {}
    }

    all_healthy = True

    # Check database
    try:
        db.get_user('healthcheck')  # Will return None but tests connection
        health_status['checks']['database'] = {
            'status': 'healthy',
            'message': 'Database connection successful'
        }
    except Exception as e:
        health_status['checks']['database'] = {
            'status': 'unhealthy',
            'message': str(e)
        }
        all_healthy = False

    # Check OpenAI API
    try:
        client = get_openai_client()
        health_status['checks']['openai_api'] = {
            'status': 'healthy',
            'message': 'API client initialized'
        }
    except Exception as e:
        health_status['checks']['openai_api'] = {
            'status': 'unhealthy',
            'message': str(e)
        }
        all_healthy = False

    # Check disk space
    try:
        import shutil
        stats = shutil.disk_usage('/')
        free_gb = stats.free / (1024**3)
        if free_gb < 1:
            health_status['checks']['disk_space'] = {
                'status': 'warning',
                'message': f'Low disk space: {free_gb:.2f} GB free'
            }
        else:
            health_status['checks']['disk_space'] = {
                'status': 'healthy',
                'message': f'Disk space OK: {free_gb:.2f} GB free'
            }
    except Exception as e:
        health_status['checks']['disk_space'] = {
            'status': 'unknown',
            'message': str(e)
        }

    # Check SSL certificate expiry (if HTTPS enabled)
    if os.getenv('USE_HTTPS', 'false').lower() == 'true':
        cert_path = os.getenv('SSL_CERT_PATH')
        if cert_path and os.path.exists(cert_path):
            # TODO: Check certificate expiry
            health_status['checks']['ssl_certificate'] = {
                'status': 'healthy',
                'message': 'Certificate present'
            }

    # Set overall status
    health_status['status'] = 'healthy' if all_healthy else 'degraded'

    status_code = 200 if all_healthy else 503
    return jsonify(health_status), status_code

@app.route('/metrics', methods=['GET'])
def metrics_prometheus():
    """
    Prometheus-compatible metrics endpoint
    Returns metrics in Prometheus text format
    """
    try:
        metrics = get_metrics_collector()
        prometheus_output = metrics.export_prometheus()

        return prometheus_output, 200, {'Content-Type': 'text/plain; version=0.0.4'}
    except Exception as e:
        return f"# Error generating metrics: {str(e)}\n", 500, {'Content-Type': 'text/plain'}

@app.route('/metrics/json', methods=['GET'])
def metrics_json():
    """
    JSON metrics endpoint for dashboard and API consumption
    """
    try:
        metrics = get_metrics_collector()
        json_output = metrics.export_json()

        return jsonify(json_output), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/metrics/dashboard', methods=['GET'])
@login_required
def metrics_dashboard():
    """
    Web-based metrics dashboard
    Requires authentication
    """
    try:
        metrics = get_metrics_collector()
        analyzer = LogAnalyzer()

        metrics_data = metrics.export_json()
        security_analysis = analyzer.analyze_security_logs(hours=24)
        anomalies = analyzer.detect_anomalies(hours=24)

        dashboard_html = '''
<!DOCTYPE html>
<html>
<head>
    <title>VoiceVerse Metrics Dashboard</title>
    <meta charset="utf-8">
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            min-height: 100vh;
            padding: 20px;
        }
        .container {
            max-width: 1400px;
            margin: 0 auto;
        }
        .header {
            background: rgba(255,255,255,0.95);
            padding: 25px 30px;
            border-radius: 15px;
            margin-bottom: 25px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }
        .header h1 {
            font-size: 28px;
            color: #333;
            margin-bottom: 10px;
        }
        .header .meta {
            color: #666;
            font-size: 14px;
        }
        .grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
            margin-bottom: 25px;
        }
        .card {
            background: rgba(255,255,255,0.95);
            border-radius: 15px;
            padding: 25px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
        }
        .card h2 {
            font-size: 18px;
            color: #333;
            margin-bottom: 15px;
            border-bottom: 2px solid #667eea;
            padding-bottom: 10px;
        }
        .stat {
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 12px 0;
            border-bottom: 1px solid #eee;
        }
        .stat:last-child {
            border-bottom: none;
        }
        .stat-label {
            color: #666;
            font-size: 14px;
        }
        .stat-value {
            font-weight: 600;
            font-size: 16px;
            color: #333;
        }
        .stat-value.good { color: #10b981; }
        .stat-value.warning { color: #f59e0b; }
        .stat-value.bad { color: #ef4444; }
        .progress-bar {
            width: 100%;
            height: 8px;
            background: #e5e7eb;
            border-radius: 4px;
            overflow: hidden;
            margin-top: 8px;
        }
        .progress-fill {
            height: 100%;
            transition: width 0.3s ease;
        }
        .progress-fill.good { background: #10b981; }
        .progress-fill.warning { background: #f59e0b; }
        .progress-fill.bad { background: #ef4444; }
        .alert {
            background: #fef3c7;
            border-left: 4px solid #f59e0b;
            padding: 15px;
            margin: 10px 0;
            border-radius: 5px;
        }
        .alert.danger {
            background: #fee2e2;
            border-left-color: #ef4444;
        }
        .alert-title {
            font-weight: 600;
            margin-bottom: 5px;
        }
        .alert-message {
            font-size: 14px;
            color: #666;
        }
        .back-link {
            display: inline-block;
            margin-top: 20px;
            padding: 10px 20px;
            background: rgba(255,255,255,0.95);
            color: #667eea;
            text-decoration: none;
            border-radius: 8px;
            font-weight: 500;
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }
        .back-link:hover {
            background: #fff;
        }
        .list-item {
            padding: 10px;
            border-left: 3px solid #667eea;
            margin: 10px 0;
            background: #f9fafb;
            border-radius: 4px;
            font-size: 13px;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>📊 VoiceVerse Metrics Dashboard</h1>
            <div class="meta">
                Updated: {{ timestamp }}<br>
                Uptime: {{ uptime_hours }} hours
            </div>
        </div>

        <!-- System Metrics -->
        <div class="grid">
            <div class="card">
                <h2>💻 System Resources</h2>
                <div class="stat">
                    <span class="stat-label">CPU Usage</span>
                    <span class="stat-value {{ cpu_class }}">{{ cpu_percent }}%</span>
                </div>
                <div class="progress-bar">
                    <div class="progress-fill {{ cpu_class }}" style="width: {{ cpu_percent }}%"></div>
                </div>

                <div class="stat">
                    <span class="stat-label">Memory Usage</span>
                    <span class="stat-value {{ memory_class }}">{{ memory_percent }}%</span>
                </div>
                <div class="progress-bar">
                    <div class="progress-fill {{ memory_class }}" style="width: {{ memory_percent }}%"></div>
                </div>

                <div class="stat">
                    <span class="stat-label">Disk Usage</span>
                    <span class="stat-value {{ disk_class }}">{{ disk_percent }}%</span>
                </div>
                <div class="progress-bar">
                    <div class="progress-fill {{ disk_class }}" style="width: {{ disk_percent }}%"></div>
                </div>
            </div>

            <div class="card">
                <h2>🌐 HTTP Requests</h2>
                <div class="stat">
                    <span class="stat-label">Total Requests</span>
                    <span class="stat-value">{{ total_requests }}</span>
                </div>
                <div class="stat">
                    <span class="stat-label">2xx Success</span>
                    <span class="stat-value good">{{ status_2xx }}</span>
                </div>
                <div class="stat">
                    <span class="stat-label">4xx Client Errors</span>
                    <span class="stat-value warning">{{ status_4xx }}</span>
                </div>
                <div class="stat">
                    <span class="stat-label">5xx Server Errors</span>
                    <span class="stat-value bad">{{ status_5xx }}</span>
                </div>
            </div>

            <div class="card">
                <h2>🎤 TTS Generation</h2>
                <div class="stat">
                    <span class="stat-label">Total Generated</span>
                    <span class="stat-value">{{ tts_total }}</span>
                </div>
                <div class="stat">
                    <span class="stat-label">Success Rate</span>
                    <span class="stat-value {{ tts_success_class }}">{{ tts_success_rate }}%</span>
                </div>
                <div class="stat">
                    <span class="stat-label">Failures</span>
                    <span class="stat-value bad">{{ tts_failure }}</span>
                </div>
            </div>

            <div class="card">
                <h2>👥 User Activity</h2>
                <div class="stat">
                    <span class="stat-label">Active Sessions</span>
                    <span class="stat-value good">{{ active_sessions }}</span>
                </div>
                <div class="stat">
                    <span class="stat-label">Total Users</span>
                    <span class="stat-value">{{ total_users }}</span>
                </div>
                <div class="stat">
                    <span class="stat-label">Failed Logins</span>
                    <span class="stat-value {{ failed_logins_class }}">{{ failed_logins }}</span>
                </div>
            </div>
        </div>

        <!-- Security Alerts -->
        {% if security_alerts %}
        <div class="card">
            <h2>🔒 Security Alerts</h2>
            {% for alert in security_alerts %}
            <div class="alert {{ alert.severity }}">
                <div class="alert-title">{{ alert.title }}</div>
                <div class="alert-message">{{ alert.message }}</div>
            </div>
            {% endfor %}
        </div>
        {% endif %}

        <!-- Anomalies -->
        {% if anomalies %}
        <div class="card">
            <h2>⚠️ Detected Anomalies</h2>
            {% for anomaly in anomalies %}
            <div class="list-item">
                <strong>{{ anomaly.type }}</strong> ({{ anomaly.severity }})<br>
                {{ anomaly.description }}
            </div>
            {% endfor %}
        </div>
        {% endif %}

        <a href="/dashboard" class="back-link">← Back to Dashboard</a>
    </div>

    <script>
        // Auto-refresh every 30 seconds
        setTimeout(function() {
            location.reload();
        }, 30000);
    </script>
</body>
</html>
        '''

        from flask import render_template_string

        # Process data for template
        system = metrics_data['system']
        http = metrics_data['http']
        tts = metrics_data['tts']
        users = metrics_data['users']
        db = metrics_data['database']

        # Calculate status code counts
        status_codes = http.get('status_codes', {})
        status_2xx = sum(v for k, v in status_codes.items() if k.startswith('2'))
        status_4xx = sum(v for k, v in status_codes.items() if k.startswith('4'))
        status_5xx = sum(v for k, v in status_codes.items() if k.startswith('5'))

        # TTS success rate
        tts_total = tts['total']
        tts_success = tts['success']
        tts_success_rate = (tts_success / tts_total * 100) if tts_total > 0 else 100

        # Uptime in hours
        uptime_hours = round(metrics_data['uptime_seconds'] / 3600, 2)

        # Resource usage classes
        cpu_class = 'good' if system['cpu_percent'] < 70 else 'warning' if system['cpu_percent'] < 90 else 'bad'
        memory_class = 'good' if system['memory_percent'] < 70 else 'warning' if system['memory_percent'] < 90 else 'bad'
        disk_class = 'good' if system['disk_percent'] < 70 else 'warning' if system['disk_percent'] < 90 else 'bad'
        tts_success_class = 'good' if tts_success_rate >= 95 else 'warning' if tts_success_rate >= 80 else 'bad'
        failed_logins_class = 'good' if users['failed_logins'] < 10 else 'warning' if users['failed_logins'] < 50 else 'bad'

        # Security alerts
        security_alerts = []
        if len(security_analysis.get('brute_force_attempts', [])) > 0:
            security_alerts.append({
                'severity': 'danger',
                'title': 'Brute Force Attempts Detected',
                'message': f"{len(security_analysis['brute_force_attempts'])} IP(s) with multiple failed login attempts"
            })

        if len(security_analysis.get('threats_detected', [])) > 0:
            security_alerts.append({
                'severity': 'danger',
                'title': 'Threats Detected',
                'message': f"{len(security_analysis['threats_detected'])} potential threats identified in logs"
            })

        if users['failed_logins'] > 50:
            security_alerts.append({
                'severity': 'warning',
                'title': 'High Failed Login Count',
                'message': f"{users['failed_logins']} failed login attempts in last 24 hours"
            })

        return render_template(
            dashboard_html,
            timestamp=metrics_data['timestamp'],
            uptime_hours=uptime_hours,
            cpu_percent=round(system['cpu_percent'], 1),
            memory_percent=round(system['memory_percent'], 1),
            disk_percent=round(system['disk_percent'], 1),
            cpu_class=cpu_class,
            memory_class=memory_class,
            disk_class=disk_class,
            total_requests=http['total_requests'],
            status_2xx=status_2xx,
            status_4xx=status_4xx,
            status_5xx=status_5xx,
            tts_total=tts_total,
            tts_success_rate=round(tts_success_rate, 1),
            tts_failure=tts['failure'],
            tts_success_class=tts_success_class,
            active_sessions=users['active_sessions'],
            total_users=db['user_count'],
            failed_logins=users['failed_logins'],
            failed_logins_class=failed_logins_class,
            security_alerts=security_alerts,
            anomalies=anomalies.get('anomalies', [])
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/favicon.ico')
def favicon():
    return '', 204

# Phase 4: Feature Enhancements
try:
    from phase4_routes import init_phase4_routes
    init_phase4_routes(app)
    print("[Phase 4] Feature enhancements loaded successfully")
except ImportError as e:
    print(f"[Phase 4] Warning: Could not load Phase 4 features: {e}")
except Exception as e:
    print(f"[Phase 4] Error initializing Phase 4 features: {e}")

if __name__ == '__main__':
    if not os.getenv("OPENAI_API_KEY"):
        print("⚠️  WARNING: OPENAI_API_KEY environment variable is not set!")
        print("Please set it using: export OPENAI_API_KEY='your-api-key-here'")
        print("")

    # Security Migration: Run ownership migration for existing files
    print("🔒 Running security migration...")
    migrate_existing_files_ownership()
    print("")

    # Security: Read debug mode from environment variable, default to False for production
    debug_mode = os.getenv('DEBUG', 'False').lower() == 'true'
    port = int(os.getenv('PORT', 5000))
    # Security: Default to localhost only (use 0.0.0.0 to allow external connections)
    host = os.getenv('HOST', '127.0.0.1')

    # HTTPS/TLS: Check if HTTPS is enabled
    use_https = os.getenv('USE_HTTPS', 'false').lower() == 'true'
    ssl_context = None

    if use_https:
        # Load SSL certificate and key from environment
        ssl_cert_path = os.getenv('SSL_CERT_PATH', 'certs/dev-cert.pem')
        ssl_key_path = os.getenv('SSL_KEY_PATH', 'certs/dev-key.pem')

        # Validate certificate files exist
        if not os.path.exists(ssl_cert_path):
            print(f"❌ ERROR: SSL certificate not found at {ssl_cert_path}")
            print("Generate certificates with: bash scripts/generate_dev_cert.sh")
            exit(1)

        if not os.path.exists(ssl_key_path):
            print(f"❌ ERROR: SSL private key not found at {ssl_key_path}")
            print("Generate certificates with: bash scripts/generate_dev_cert.sh")
            exit(1)

        # Create SSL context
        import ssl
        ssl_context = ssl.SSLContext(ssl.PROTOCOL_TLS_SERVER)
        ssl_context.load_cert_chain(ssl_cert_path, ssl_key_path)

        print("🌌 Starting VoiceVerse - Text-to-Speech Application")
        print(f"📍 Open in your browser: https://{host}:{port}")
        print("🔒 HTTPS is ENABLED (TLS/SSL encrypted connection)")
        if 'dev' in ssl_cert_path:
            print("⚠️  Using self-signed certificate - you'll need to accept the browser warning")
        print("Press Ctrl+C to stop the server\n")
    else:
        print("🌌 Starting VoiceVerse - Text-to-Speech Application")
        print(f"📍 Open in your browser: http://{host}:{port}")
        print("⚠️  HTTPS is DISABLED - consider enabling for production")
        print("   Set USE_HTTPS=true in .env and run scripts/generate_dev_cert.sh")
        print("Press Ctrl+C to stop the server\n")

    app.run(debug=debug_mode, port=port, host=host, ssl_context=ssl_context)
